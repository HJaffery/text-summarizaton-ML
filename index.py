import docx
import os, sys, email
import numpy as np
import pandas as pd
import re

#takes the file name which needs to be summarized
filename = ''
filename = input('Input filename:')
filename = filename + '.docx'

# creates a Document to read the input file
doc = docx.Document(filename)

text = ""

# getting the text form input file to a string
for para in doc.paragraphs:
    text += para.text


#splits text into sentences
from nltk.tokenize import sent_tokenize
x = sent_tokenize(text)

#giving path to use the pre buit skipthoughts netwrok
#skipsthoughts is taken form https://github.com/ryankiros/skip-thoughts
import sys
sys.path.insert(0, 'F:\Data Science\EML\Text summarization\skip-thoughts-master')
import skipthoughts

print('loading skip thoughts model')
model = skipthoughts.load_model()

print('Making encoder')
encoder = skipthoughts.Encoder(model)


print('making vectors')
vectors = encoder.encode(x)
print (vectors)

# putting vectors to a csv file. Just to explore or see vectors
import pandas as pd
vectorsdf = pd.DataFrame(vectors)
vectorsdf.to_csv('kiterunner_vectors.csv')

import numpy as np
from sklearn.cluster import KMeans

print('clustering')
n_clusters = np.ceil(len(vectors)**0.7) # the factor 0.7 could be decreased or increased to reduce or increase the numbers of sentences included in summary
                                        # above line is just simple mathematics to readuce text to a certain factor

# putting vectors to different clusters
n_clusters = int(n_clusters)
kmeans = KMeans(n_clusters=n_clusters)
kmeans = kmeans.fit(vectors)

from sklearn.metrics import pairwise_distances_argmin_min
# finding the most sutiable candidate sentence form each cluster by using the fact that the sentence nearest to  cluster mean is representative of the cluster
# and is most eligible to be selected to be a part of the summary

print('creating summary')
avg = []
for j in range(n_clusters):
    idx = np.where(kmeans.labels_ == j)[0]
    avg.append(np.mean(idx))
closest, _ = pairwise_distances_argmin_min(kmeans.cluster_centers_, vectors)
ordering = sorted(range(n_clusters), key=lambda k: avg[k])
#print(ordering)
summary = ' '.join([x[closest[idx]] for idx in ordering])

print(summary)

# saving the summary to a .docx file
doc = docx.Document()
doc.add_paragraph(summary)
doc.save('kiterunner_summary.docx')

